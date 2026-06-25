import csv, sys
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("Run: pip install python-docx --break-system-packages")

HERE = Path(__file__).parent
doc = Document()
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(10)

def caption(n, text):
    p = doc.add_paragraph()
    r = p.add_run(f"Table {n}. {text}")
    r.bold = True; r.font.size = Pt(10)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i].paragraphs[0].add_run(h); c.bold = True; c.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i,v in enumerate(row):
            cells[i].paragraphs[0].add_run(str(v)).font.size = Pt(9)
    doc.add_paragraph()

# ---------- TABLE 1: community summary ----------
LABELS = {0:"Performance & Productivity",1:"Work-Life / Teleworking",2:"Belonging & Workplace",
          3:"Technostress & Digital Fatigue",4:"Psychological Safety & Leadership",5:"Surveillance & Monitoring"}
# ordinal orientation (NOT cardinal nets - per the stability finding)
ORIENT = {0:"benefit-leaning",1:"near-neutral*",2:"benefit-leaning",3:"near-neutral*",
          4:"most benefit-leaning",5:"most risk-leaning"}
t1=[]
ct = HERE/"rq1_analysis"/"community_table.csv"
sizes={}
for r in csv.DictReader(open(ct)):
    c=int(r["community"]); sizes[c]=int(r["size"])
for c in sorted(LABELS):
    t1.append([c, LABELS[c], sizes.get(c,"?"), ORIENT[c]])
caption(1, "Thematic communities recovered by Leiden detection (n = 387). Orientation reported ordinally; mid-field communities (*) are partition-sensitive and not robust point estimates (ARI = 0.78).")
add_table(["Community","Theme","n papers","Orientation"], t1)

# ---------- TABLE 2: construct centrality ----------
TARGETS = ["trust","psychological safety","surveillance","autonomy","belonging","technostress"]
cc = HERE/"rq3_constructs"/"concept_centrality.csv"
rows2=[]
for r in csv.DictReader(open(cc)):
    if r["concept"] in TARGETS:
        rows2.append([r["concept"], r["degree"], r["betweenness"], r["eigenvector"],
                      r["rank_degree"], r["rank_betweenness"], r["composite_rank"]])
# order by composite rank
rows2.sort(key=lambda x: int(x[6]))
caption(2, "Centrality of target psychosocial constructs in the concept co-occurrence network (RQ3). Autonomy brokers the field; trust is pervasive; surveillance is peripheral.")
add_table(["Construct","Degree","Betweenness","Eigenvector","Rank(deg)","Rank(btw)","Composite rank"], rows2)

# ---------- TABLE 3: benefit/risk co-presence ----------
cd = HERE/"rq4_prevalence"/"copresence_distribution.csv"
rows3=[]
for r in csv.DictReader(open(cd)):
    rows3.append([f">={r['threshold']}", f"{r['pct_both']}%", f"{r['pct_benefit_only']}%",
                  f"{r['pct_risk_only']}%", f"{r['pct_neither']}%"])
caption(3, "Per-paper benefit/risk co-presence at anchor thresholds (RQ4; seed-independent). The field is polarized at the community level but individual papers engage both poles.")
add_table(["Anchors/pole","Both","Benefit only","Risk only","Neither"], rows3)

# ---------- TABLE 4: outcome measurement by level (+validation) ----------
me = HERE/"measured_extract"/"measured_outcomes.csv"
from collections import Counter
lvl = Counter()
total_dv = 0
for r in csv.DictReader(open(me)):
    if r["role"]=="measured_DV" and r["level"] in ("individual","team","firm"):
        lvl[r["level"]]+=1; total_dv+=1
rows4=[]
for L in ["individual","team","firm"]:
    n=lvl[L]; pct = round(100*n/total_dv,1) if total_dv else 0
    rows4.append([L.capitalize(), n, f"{pct}%"])
rows4.append(["Total measured DVs", total_dv, "100%"])
caption(4, "Level of analysis at which outcomes are measured as dependent variables (LLM extraction, n = 387; Claude claude-sonnet-4-6, Anthropic). Validation against a blind independent coder: level agreement 81-86%, firm-level precision 78% (18/23 confirmed). The measured-versus-mentioned distinction was only moderately reliable (kappa = 0.48) and is reported in text as suggestive only.")
add_table(["Measurement level","Measured-DV instances","Share"], rows4)

out = HERE/"manuscript_tables.docx"
doc.save(out)
print(f"Saved {out}")
print(f"Table 1: {len(t1)} communities | Table 2: {len(rows2)} constructs | Table 3: {len(rows3)} thresholds | Table 4: {total_dv} measured DVs by level")
